import os
import re
import sys
import subprocess
import tkinter as tk
import threading
import time
import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
		
from core.wisper import run_wisper
from core.translate import run_translate as translate_file
from auto_compile import auto_compile_project
from core.align_author_text import process, srt_to_internal_format
from core.yt_dlp_helper import run_command, build_commands

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

ENGINE_FILE = os.path.join(os.path.dirname(__file__), "main.py")

# ======================================================
# OUTPUT SYSTEM (UNIFIED)
# ======================================================

def sanitize_filename(name: str) -> str:
    return "".join(c for c in name if c not in r'\/:*?"<>|').strip()

def make_name(stage: str, input_path_or_text: str = None, custom: str = None):
    if custom:
        base = custom

    elif input_path_or_text:
        if input_path_or_text.startswith("http"):
            base = "youtube_video"
        else:
            base = os.path.splitext(
                os.path.basename(input_path_or_text)
            )[0]

            base = re.sub(r'^[A-Z_]+_', '', base)

    else:
        base = "project"

    stage = sanitize_filename(stage)
    base = sanitize_filename(base)

    return f"{stage}_{base}" if stage else base

def build_output_path(input_path, folder, stage, custom_name, default_ext) -> str:
    """
    Формат:
        folder / STAGE_project.ext
    """

    name = make_name(stage, input_path, custom_name)
    return os.path.join(folder, f"{name}{default_ext}")

# ======================================================
# LANGUAGES
# ======================================================

LANGUAGES = {
    "uk": {
        "title": "Інструмент субтитрів — Візуальна версія",
        "tab_translate": "Переклад SRT",
        "tab_compile": "Збірка автора",
        "tab_wisper": "Wisper аудіо",
        "tab_yt": "YouTube Downloader",
        "tab_align": "Алігн авторського тексту",

        "choose_srt": "Обрати SRT файл",
        "run_translate": "Запустити переклад",

        "original_srt": "Оригінальний SRT",
        "author_doc": "DOCX автора",
        "run_compile": "Запустити збірку",

        "choose_audio": "Обрати аудіофайл",
        "run_wisper": "Запустити транскрипцію",

        "choose_original": "Оригінальний SRT",
        "choose_author": "Текст автора",
        "run_align": "Запустити вирівнювання",

        "error": "Помилка",
        "result": "Результат",
        "choose_file": "Оберіть файл",
        "choose_two_files": "Оберіть обидва файли",
        "choose_folder": "Обрати папку",

        "processing": "Обробка...",
        "done": "Готово",

        "lang_ua": "Українська",
        "lang_en": "English"
    },

    "en": {
        "title": "Subtitle Tool — Visual Edition",
        "tab_translate": "Translate SRT",
        "tab_compile": "Compile Author",
        "tab_wisper": "Wisper Audio",
        "tab_yt": "YouTube Downloader",
        "tab_align": "Align Author Text",

        "choose_srt": "Choose SRT file",
        "run_translate": "Run Translation",

        "original_srt": "Original SRT",
        "author_doc": "Author DOCX",
        "run_compile": "Run Compile",

        "choose_audio": "Choose Audio",
        "run_wisper": "Run Transcription",

        "choose_original": "Original SRT",
        "choose_author": "Author Text",
        "run_align": "Run Align",

        "error": "Error",
        "result": "Result",
        "choose_file": "Choose file",
        "choose_two_files": "Choose both files",
        "choose_folder": "Choose folder",
        "processing": "Processing...",
        "done": "Done",

        "lang_ua": "Українська",
        "lang_en": "English"
    }
}


# ======================================================
# APP
# ======================================================

class App(ctk.CTk):

    def __init__(self, lang="uk"):
        super().__init__()

        self.lang = lang
        self.translations = LANGUAGES[self.lang]

        self.title(self.t("title"))
        self.geometry("900x600")

        # OUTPUT defaults
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        self.output_folder = ctk.StringVar(value=desktop)

        self.bind_all("<KeyPress>", self.handle_global_keys)

        self.create_language_switch()

        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(fill="both", expand=True, padx=20, pady=60)

        self.build_tabs()

        # =========================
        # PROGRESS BAR (GLOBAL)
        # =========================
        self.progress = ctk.CTkProgressBar(self)
        self.progress.pack(fill="x", padx=20, pady=(0, 2))
        self.progress.set(0)

        self.status_global = ctk.CTkLabel(self, text="", anchor="center")
        self.status_global.pack(fill="x", padx=20, pady=(0, 8))

    # ======================================================

    def create_progress_window(self, title="Processing..."):
        self.progress_window = ctk.CTkToplevel(self)
        self.progress_window.title(title)
        self.progress_window.geometry("400x120")
        self.progress_window.resizable(False, False)

        self.progress_window.grab_set()

        self.progress_label = ctk.CTkLabel(
            self.progress_window,
            text=title
        )
        self.progress_label.pack(pady=(20, 10))

        self.progress_bar_popup = ctk.CTkProgressBar(
            self.progress_window,
            width=320
        )
        self.progress_bar_popup.pack(pady=10)
        self.progress_bar_popup.set(0)

    def update_yt_progress(self, percent):
        value = percent / 100

        def _update():
            self.progress.set(value)
            self.status_global.configure(
                text=f"Downloading... {percent:.1f}%"
            )

            if hasattr(self, "progress_bar_popup"):
                self.progress_bar_popup.set(value)

            if hasattr(self, "progress_label"):
                self.progress_label.configure(
                    text=f"Downloading... {percent:.1f}%"
                )

            if percent >= 100:
                self.set_status("Finalizing...")

        self.after(0, _update)

    def update_progress_translate(self, percent):
        def _update():
            value = percent / 100
            self.progress.set(value)
            self.status_global.configure(
                text=f"Translating... {percent:.1f}%"
            )
        self.after(0, _update)

    # ======================================================

    def run_command(self, cmd, progress_callback=None):
        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            encoding="utf-8",
            errors="ignore"
        )

        for line in process.stdout:
            print(line, end="")

            match = re.search(r'(\d+(?:\.\d+)?)%', line)

            if match and progress_callback:
                percent = float(match.group(1))
                self.after(
                    0,
                    lambda p=percent: progress_callback(p)
                )

        process.wait()
        return process.returncode


    def update_progress_window(self, current, total, text=""):
        if not hasattr(self, "progress_window"):
            return

        percent = current / total

        def _update():
            self.progress_label.configure(
                text=f"{text} ({int(percent*100)}%)"
            )
            self.progress_bar_popup.set(percent)

        self.after(0, _update)
    

    def close_progress_window(self):
        if hasattr(self, "progress_window"):
            self.after(
                0,
                lambda: self.progress_window.destroy()
            )


    def start_progress(self, determinate=True):
        def _start():
            if determinate:
                self.progress.stop()
                self.progress.configure(mode="determinate")
            else:
                self.progress.configure(mode="indeterminate")
                self.progress.start()

        self.after(0, _start)

    def stop_progress(self):
        def _stop():
            self.progress.stop()
            self.progress.configure(mode="determinate")
            self.progress.set(0)
        self.after(0, _stop)

    def progress_step(self, current, total, label=""):
        if total == 0:
            return

        percent = current / total
        text = f"{label} {current}/{total} ({int(percent * 100)}%)"

        self.set_status(text, percent)
        self.update_progress_window(current, total, label)

    def set_status(self, text: str, progress: float = None):
        def _update():
            self.status_global.configure(text=text)

            if progress is not None:
                self.progress.set(progress)

        self.after(0, _update)

    # ======================================================

    def select_all(self, widget):
        try:
            widget.select_range(0, tk.END)
            widget.icursor(tk.END)
        except Exception:
            pass


    def copy_selection(self, widget):
        try:
            selected = widget.selection_get()
            self.clipboard_clear()
            self.clipboard_append(selected)
        except Exception:
            pass


    def cut_selection(self, widget):
        try:
            selected = widget.selection_get()
            self.clipboard_clear()
            self.clipboard_append(selected)

            widget.delete("sel.first", "sel.last")
        except Exception:
            pass

    # ======================================================

    def handle_global_keys(self, event):
        ctrl = (event.state & 0x4) != 0
        widget = event.widget

        # Вставка (Ctrl+V)
        if ctrl and event.keycode == 86:
            self.paste_to_entry(widget)
            return "break"

        # Виділити все (Ctrl+A)
        if ctrl and event.keycode == 65:
            self.select_all(widget)
            return "break"

        # Копіювати (Ctrl+C)
        if ctrl and event.keycode == 67:
            self.copy_selection(widget)
            return "break"

        # Вирізати (Ctrl+X)
        if ctrl and event.keycode == 88:
            self.cut_selection(widget)
            return "break"

    # ======================================================

    def show_left_click_menu(self, event):
        widget = event.widget

        try:
            widget.focus_set()
        except:
            return

        # Створюємо меню
        menu = tk.Menu(self, tearoff=0)
        menu.add_command(
            label="Paste",
            command=lambda: self.paste_to_entry(widget)
        )

        self.after(50, lambda: menu.tk_popup(event.x_root, event.y_root))

    # ======================================================

    def reset_ui_state(self):
        if hasattr(self, "btn_translate"):
            self.btn_translate.configure(state="normal")

        if hasattr(self, "status_translate"):
            self.status_translate.configure(text=self.t("done"))

    # ======================================================

    def create_output_block(self, parent):
        frame = ctk.CTkFrame(parent)
        frame.pack(fill="x", padx=10, pady=5)

        frame.grid_columnconfigure(0, weight=1)

        # поле папки
        entry = ctk.CTkEntry(
            frame,
            textvariable=self.output_folder,
            height=36
        )
        entry.grid(
            row=0,
            column=0,
            sticky="ew",
            padx=(10, 10),
            pady=10
        )

        # кнопка вибору папки
        btn = ctk.CTkButton(
            frame,
            text=self.t("choose_folder"),
            width=120,
            command=lambda: self.output_folder.set(
                self.pick_folder()
            )
        )
        btn.grid(
            row=0,
            column=1,
            padx=(0, 10),
            pady=10
        )

        self.bind_paste_support(entry)

# ==========================================
# UNIVERSAL PASTE SYSTEM
# ==========================================

    def bind_paste_support(self, widget):
        widget.bind("<Shift-Insert>", self.handle_paste)
        widget.bind("<Button-3>", self.show_paste_menu)

        widget.bind("<Button-1>", self.show_left_click_menu)

    def handle_paste(self, event):
        self.paste_to_entry(event.widget)
        return "break"

    def show_paste_menu(self, event):
        menu = tk.Menu(self, tearoff=0)
        menu.add_command(
            label="Paste",
            command=lambda: self.paste_to_entry(event.widget)
        )
        menu.tk_popup(event.x_root, event.y_root)

    # =========================
    # HELPERS
    # =========================

    def t(self, key):
        return self.translations.get(key, key)

    def error(self, msg):
        messagebox.showerror(self.t("error"), msg)

    def info(self, msg):
        messagebox.showinfo(self.t("result"), msg)

    def run_async(self, func, title="Processing..."):
        def wrapper():
            try:
                self.after(0, lambda: self.set_status(title))
                self.start_progress(determinate=True)

                func()

            except Exception as e:
                err = str(e)
                self.after(0, lambda err=err: self.error(err))

            finally:
                self.after(0, self.stop_progress)
                self.after(0, lambda: self.set_status("Done"))
                self.after(0, self.reset_ui_state)

        t = threading.Thread(target=wrapper, daemon=True)
        t.start()

    def pick_file(self, types):
        return filedialog.askopenfilename(filetypes=types)

    def pick_folder(self):
        return filedialog.askdirectory()

    # =========================
    # LANGUAGE
    # =========================

    def create_language_switch(self):
        self.lang_menu = ctk.CTkOptionMenu(
            self,
            values=[self.t("lang_ua"), self.t("lang_en")],
            command=self.change_language
        )
        self.lang_menu.set(self.t("lang_ua"))
        self.lang_menu.place(relx=1.0, x=-20, y=20, anchor="ne")

    def change_language(self, choice):
        self.lang = "uk" if "Україн" in choice else "en"
        self.translations = LANGUAGES[self.lang]
        self.title(self.t("title"))

        self.tabview.destroy()
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(fill="both", expand=True, padx=20, pady=60)

        self.build_tabs()

    # =========================
    # TABS
    # =========================

    def build_tabs(self):
        self.tab_yt = self.tabview.add(self.t("tab_yt"))
        self.tab_wisper = self.tabview.add(self.t("tab_wisper"))
        self.tab_translate = self.tabview.add(self.t("tab_translate"))
        self.tab_align = self.tabview.add(self.t("tab_align"))
        self.tab_compile = self.tabview.add(self.t("tab_compile"))

        self.create_yt_tab()
        self.create_wisper_tab()
        self.create_translate_tab()
        self.create_align_tab()
        self.create_compile_tab()


    # ======================================================
    # RUNS
    # ======================================================

    def run_translate(self):
        if not self.translate_path.get():
            return self.error(self.t("choose_file"))

        if self.btn_translate.cget("state") == "disabled":
            return

        self.btn_translate.configure(state="disabled")
        self.status_translate.configure(text=self.t("processing"))

        input_file = self.translate_path.get()
        name = "UK"

        def task():
            try:
                self.set_status("Preparing translation...")

                # fallback прогрес
                fake_progress = {"value": 0}

                def safe_progress(percent):
                
                    fake_progress["value"] = percent
                    self.update_progress_translate(percent)

                def fake_progress_runner():
                    last = 0
                    while fake_progress["value"] < 100:
                        last = min(last + 1, 95)
                        self.update_progress_translate(last)
                        time.sleep(0.2)

                watcher = threading.Thread(target=fake_progress_runner, daemon=True)
                watcher.start()

                result = translate_file(
                    input_file,
                    target_lang="uk",
                    output_folder=self.output_folder.get(),
                    output_name=name,
                    progress_callback=safe_progress
                )

                fake_progress["value"] = 100
                self.update_progress_translate(100)

                folder = result.get("folder", "—")

                self.after(0, lambda: self.info("Done"))

            finally:
                self.after(0, self.on_translate_done)

        self.run_async(task, "Translating...")

    def on_translate_done(self):
        self.status_translate.configure(text=self.t("done"))
        self.btn_translate.configure(state="normal")

    # ======================================================

    def run_compile(self):
        if not self.compile_srt.get() or not self.compile_doc.get():
            return self.error(self.t("choose_two_files"))

        def task():
            res = auto_compile_project(
                self.compile_srt.get(),
                self.compile_doc.get()
            )

            total = len(res)

            blocks = []
            for i, (line, start, end) in enumerate(res, start=1):
                self.progress_step(i, total, "Compile SRT")
                blocks.append(
                    f"{i}\n"
                    f"{start} --> {end}\n"
                    f"{line}\n"
                )

            text = "\n".join(blocks).strip() + "\n"

            name = make_name("MY", self.compile_srt.get())

            output_path = os.path.join(
                self.output_folder.get(),
                name + ".srt"
            )

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)

            self.after(0, lambda: self.info(f"Saved: {output_path}"))

        self.run_async(task, "Compiling...")

    # ======================================================

    def run_wisper(self):
        if not self.audio_path.get():
            return self.error(self.t("choose_file"))
        
        self.after(0, self.stop_progress)
        self.after(0, lambda: self.set_status(self.t("done")))
        self.set_status("Processing audio...")

        def task():
            audio = self.audio_path.get()

            stage = f"W_{self.model.get().upper()}"
            name = make_name(stage, audio)

            output_path = build_output_path(
                audio,
                self.output_folder.get(),
                stage,
                None,
                ".srt"
            )

            run_wisper(
                audio,
                model_name=self.model.get(),
                output_path=output_path
            )

            self.after(0, lambda: self.set_status(self.t("done")))

        self.run_async(task, "Loading Whisper...")

    # ======================================================

    def run_yt(self):
        if not self.yt_url.get():
            return self.error("Enter URL")

        def task():
            try:
                mode_key = self.download_mode.get()
                mode = self.mode_map.get(mode_key, "reaper")

                cmds = build_commands(
                    self.yt_url.get(),
                    mode=mode,
                    sub_lang="en",
                    output_folder=self.output_folder.get(),
                )

                total = len(cmds)

                for i, cmd in enumerate(cmds, start=1):

                    # постпроцесинг ffmpeg
                    if isinstance(cmd, tuple) and cmd[0] == "POSTPROCESS":
                        _, func, folder = cmd
                        ffmpeg_cmd = func(folder)

                        if ffmpeg_cmd:
                            self.set_status("Postprocessing...")
                            self.run_command(ffmpeg_cmd)

                    else:
                        self.set_status(f"Downloading step {i}/{total}")

                        self.run_command(
                            cmd,
                            progress_callback=self.update_yt_progress
                        )

                self.after(
                    0,
                    lambda: self.info("Download finished")
                )

            except Exception as e:
                self.after(
                    0,
                    lambda: self.error(str(e))
                )

        self.run_async(task, "Downloading...")

        self.after(0, self.close_progress_window)

    # ======================================================

    def read_docx(self, path):
        doc = Document(path)
        lines = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)

        return "\n".join(lines)

    # ======================================================

    def run_align(self):
        if not self.align_srt.get() or not self.align_doc.get():
            return self.error(self.t("choose_two_files"))

        def task():
            with open(self.align_srt.get(), encoding="utf-8") as f:
                original = srt_to_internal_format(f.read())

            if self.align_doc.get().endswith(".docx"):
                doc = Document(self.align_doc.get())
                author = "\n".join(p.text for p in doc.paragraphs)
            else:
                with open(self.align_doc.get(), encoding="utf-8") as f:
                    author = f.read()

            result = process(original, author)

            name = make_name("ALI", self.align_srt.get())

            path = os.path.join(
                self.output_folder.get(),
                name + ".docx"
            )

            self.save_to_docx(result, path)

            self.after(0, lambda: self.info(f"Saved: {path}"))

        self.run_async(task, "Aligning...")

    # ======================================================
    # TRANSLATE
    # ======================================================

    def create_translate_tab(self):
        self.translate_path = ctk.StringVar()

        # INPUT
        input_frame = ctk.CTkFrame(self.tab_translate)
        input_frame.pack(fill="x", padx=15, pady=(15, 8))
        input_frame.grid_columnconfigure(0, weight=1)

        entry = ctk.CTkEntry(
            input_frame,
            textvariable=self.translate_path,
            height=38
        )
        entry.grid(
            row=0,
            column=0,
            sticky="ew",
            padx=(10, 8),
            pady=10
        )

        ctk.CTkButton(
            input_frame,
            text=self.t("choose_srt"),
            width=120,
            command=lambda: self.translate_path.set(
                self.pick_file([("SRT", "*.srt")])
            )
        ).grid(
            row=0,
            column=1,
            padx=(0, 10),
            pady=10
        )

        # STATUS
        status_frame = ctk.CTkFrame(
            self.tab_translate,
            fg_color="transparent"
        )
        status_frame.pack(fill="x", padx=15, pady=(5, 5))

        self.status_translate = ctk.CTkLabel(
            status_frame,
            text=""
        )
        self.status_translate.pack(anchor="center")

        # ACTION
        self.btn_translate = ctk.CTkButton(
            self.tab_translate,
            text=self.t("run_translate"),
            command=self.run_translate,
            height=42
        )
        self.btn_translate.pack(
            fill="x",
            padx=15,
            pady=(10, 15)
        )

        self.create_output_block(self.tab_translate)

    # ======================================================
    # COMPILE
    # ======================================================

    def create_compile_tab(self):
        self.compile_srt = ctk.StringVar()
        self.compile_doc = ctk.StringVar()

        # SRT
        frame1 = ctk.CTkFrame(self.tab_compile)
        frame1.pack(fill="x", padx=15, pady=(15, 6))
        frame1.grid_columnconfigure(0, weight=1)

        ctk.CTkEntry(
            frame1,
            textvariable=self.compile_srt,
            height=38
        ).grid(
            row=0,
            column=0,
            sticky="ew",
            padx=(10, 8),
            pady=6
        )

        ctk.CTkButton(
            frame1,
            text=self.t("original_srt"),
            width=120,
            command=lambda: self.compile_srt.set(
                self.pick_file([("SRT", "*.srt")])
            )
        ).grid(
            row=0,
            column=1,
            padx=(0, 10),
            pady=6
        )

        # DOC
        frame2 = ctk.CTkFrame(self.tab_compile)
        frame2.pack(fill="x", padx=15, pady=(0, 8))
        frame2.grid_columnconfigure(0, weight=1)

        ctk.CTkEntry(
            frame2,
            textvariable=self.compile_doc,
            height=38
        ).grid(
            row=0,
            column=0,
            sticky="ew",
            padx=(10, 8),
            pady=6
        )

        ctk.CTkButton(
            frame2,
            text=self.t("author_doc"),
            width=120,
            command=lambda: self.compile_doc.set(
                self.pick_file([("DOCX", "*.docx"), ("TXT", "*.txt")])
            )
        ).grid(
            row=0,
            column=1,
            padx=(0, 10),
            pady=6
        )

        # ACTION
        ctk.CTkButton(
            self.tab_compile,
            text=self.t("run_compile"),
            command=self.run_compile,
            height=42
        ).pack(
            fill="x",
            padx=15,
            pady=(10, 15)
        )

        self.create_output_block(self.tab_compile)

    # ======================================================
    # WISPER
    # ======================================================

    def create_wisper_tab(self):
        self.audio_path = ctk.StringVar()
        self.model = ctk.StringVar(value="small")

        # =========================
        # INPUT ROW
        # =========================
        input_frame = ctk.CTkFrame(self.tab_wisper)
        input_frame.pack(fill="x", padx=15, pady=(15, 8))
        input_frame.grid_columnconfigure(0, weight=1)

        entry = ctk.CTkEntry(
            input_frame,
            textvariable=self.audio_path,
            height=38
        )
        entry.grid(row=0, column=0, sticky="ew", padx=(10, 8), pady=10)

        ctk.CTkButton(
            input_frame,
            text=self.t("choose_audio"),
            width=120,
            command=lambda: self.audio_path.set(
                self.pick_file([("Audio", "*.mp3 *.wav *.flac *.mp4 *.m4a")])
            )
        ).grid(row=0, column=1, padx=(0, 10), pady=10)

        # =========================
        # OPTIONS ROW (CENTERED)
        # =========================
        options_frame = ctk.CTkFrame(self.tab_wisper)
        options_frame.pack(fill="x", padx=15, pady=5)

        # центр-обгортка
        center_frame = ctk.CTkFrame(options_frame, fg_color="transparent")
        center_frame.pack(expand=True)

        ctk.CTkLabel(center_frame, text="Model:").pack(
            side="left", padx=(0, 8), pady=10
        )

        ctk.CTkOptionMenu(
            center_frame,
            values=["tiny", "base", "small", "medium", "large"],
            variable=self.model,
            width=140
        ).pack(side="left", pady=10)

        # =========================
        # STATUS (CENTERED)
        # =========================
        status_frame = ctk.CTkFrame(self.tab_wisper, fg_color="transparent")
        status_frame.pack(fill="x", padx=15, pady=(5, 5))

        self.status = ctk.CTkLabel(
            status_frame,
            text="Ready"
        )
        self.status.pack(anchor="center")

        # =========================
        # ACTION
        # =========================
        ctk.CTkButton(
            self.tab_wisper,
            text=self.t("run_wisper"),
            command=self.run_wisper,
            height=42
        ).pack(fill="x", padx=15, pady=(10, 15))

        # =========================
        # OUTPUT
        # =========================
        self.create_output_block(self.tab_wisper)

    # ======================================================
    # YT
    # ======================================================

    def create_yt_tab(self):
        self.yt_url = ctk.StringVar()

        self.mode_map = {
            "REAPER": "reaper",
            "Video (MP4)": "video",
            "Audio (MP3)": "audio",
            # "Subtitles": "subs"
        }

        self.download_mode = ctk.StringVar(value="REAPER")

        input_frame = ctk.CTkFrame(self.tab_yt)
        input_frame.pack(fill="x", padx=15, pady=(15, 8))
        input_frame.grid_columnconfigure(0, weight=1)

        ctk.CTkEntry(
            input_frame,
            textvariable=self.yt_url,
            height=38
        ).grid(row=0, column=0, sticky="ew", padx=(10, 8), pady=10)

        # MODE
        mode_frame = ctk.CTkFrame(self.tab_yt)
        mode_frame.pack(fill="x", padx=15, pady=5)

        center_mode = ctk.CTkFrame(mode_frame, fg_color="transparent")
        center_mode.pack(expand=True)

        ctk.CTkLabel(center_mode, text="Mode:").pack(side="left", padx=(0, 8))

        self.mode_option = ctk.CTkOptionMenu(
            center_mode,
            values=list(self.mode_map.keys()),
            variable=self.download_mode,
            width=220
        )
        self.mode_option.pack(side="left")

        ctk.CTkButton(
            self.tab_yt,
            text="Download",
            command=self.run_yt,
            height=42
        ).pack(fill="x", padx=15, pady=(10, 15))

        self.create_output_block(self.tab_yt)

    # ======================================================
    # ALIGN
    # ======================================================

    def create_align_tab(self):
        self.align_srt = ctk.StringVar()
        self.align_doc = ctk.StringVar()

        for i, (var, text) in enumerate([
            (self.align_srt, self.t("choose_original")),
            (self.align_doc, self.t("choose_author"))
        ]):
            frame = ctk.CTkFrame(self.tab_align)

            frame.pack(
                fill="x",
                padx=15,
                pady=(15, 6) if i == 0 else (0, 6)
            )

            frame.grid_columnconfigure(0, weight=1)

            ctk.CTkEntry(
                frame,
                textvariable=var,
                height=38
            ).grid(
                row=0,
                column=0,
                sticky="ew",
                padx=(10, 8),
                pady=6
            )

            ctk.CTkButton(
                frame,
                text=text,
                width=120,
                command=lambda v=var:
                    v.set(self.pick_file([("All", "*.*")]))
            ).grid(
                row=0,
                column=1,
                padx=(0, 10),
                pady=6
            )

        ctk.CTkButton(
            self.tab_align,
            text=self.t("run_align"),
            command=self.run_align,
            height=42
        ).pack(
            fill="x",
            padx=15,
            pady=(10, 15)
        )

        self.create_output_block(self.tab_align)

    # ======================================================
    # DOCX
    # ======================================================

    def save_to_docx(self, text, path):
        doc = Document()
        doc.add_heading("Aligned Subtitles", 0)

        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.save(path)

# ==========================================
# ВСТАВКА ТЕКСТУ (Ctrl+V)
# ==========================================
    def paste_to_entry(self, entry):
        try:
            text = self.clipboard_get()
            try:
                start = entry.index("sel.first")
                end = entry.index("sel.last")
                entry.delete(start, end)
            except tk.TclError:
                pass

            pos = entry.index("insert")
            entry.insert(pos, text)
        except tk.TclError:
            pass

    def handle_paste(self, event):
        self.paste_to_entry(event.widget)
        return "break"

# ======================================================
# START
# ======================================================

if __name__ == "__main__":
    App().mainloop()