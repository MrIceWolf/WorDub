import os
import re
import json
from deep_translator import GoogleTranslator
from docx import Document

# ==========================================================
# =====================  TRANSLATE  ========================
# ==========================================================

def parse_srt_blocks(content: str):
    return re.split(r"\r?\n\r?\n+", content.strip())

def extract_timing(lines):
    for line in lines:
        if "-->" in line:
            return line.strip()
    return None

def run_translate(
    input_path,
    target_lang="uk",
    output_folder=None,
    output_name=None,
    progress_callback=None,
    log_callback=None
):

    if log_callback:
        log_callback("Subtitle AutoLocalization v3.0 started")

    # =============================
    # НАЛАШТУВАННЯ
    # =============================

    MAX_CHARS = 1200
    CACHE_FILE = os.path.join(os.path.dirname(__file__), "translation_cache.json")

    translator = GoogleTranslator(source="auto", target=target_lang)

    # =============================
    # ПЕРЕВІРКА ФАЙЛУ
    # =============================

    if not os.path.isfile(input_path):
        raise FileNotFoundError(f"Файл не знайдено: {input_path}")

    # =============================
    # КЕШ
    # =============================

    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r", encoding="utf-8") as f:
                cache = json.load(f)
        except:
            cache = {}
    else:
        cache = {}

    def save_cache():
        with open(CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)

    def safe_translate(text):
        if text in cache:
            return cache[text]

        try:
            result = translator.translate(text)

            # якщо переклад порожній
            if not result or result.strip() == "":
                result = text

            # якщо переклад == оригіналу -> пробуємо ще раз у lower-case
            elif result.strip() == text.strip():
                retry = translator.translate(text.lower())
                if retry and retry.strip() != text.strip():
                    result = retry

        except Exception:
            result = text

        cache[text] = result
        save_cache()
        return result

    # =============================
    # ЧИТАННЯ SRT
    # =============================

    if log_callback:
        log_callback("Reading SRT file...")

    with open(input_path, "r", encoding="utf-8-sig") as f:
        content = f.read()

    blocks = parse_srt_blocks(content)

    subtitle_entries = []

    for block in blocks:
        lines = block.strip().split("\n")

        timing = extract_timing(lines)

        if timing:
            text_lines = [l for l in lines if l != timing and not l.strip().isdigit()]
            text = "\n".join(text_lines)
            subtitle_entries.append((timing, text))
        else:
            subtitle_entries.append((None, block))

    total_blocks = sum(1 for t, _ in subtitle_entries if t is not None)

    if log_callback:
        log_callback(f"Found {total_blocks} subtitle lines")
    print(" Починається переклад...\n")

    # =============================
    # CHUNK-ПЕРЕКЛАД
    # =============================

    translated_entries = [None] * len(subtitle_entries)

    bar_length = 30
    processed = [0]

    def update_progress(processed, total_blocks):
        percent = int((processed[0] / total_blocks) * 100)
        if progress_callback:
            progress_callback(percent)
        progress = int((processed[0] / total_blocks) * bar_length)
        bar = "█" * progress + "-" * (bar_length - progress)
        print(f" [{bar}] {percent}% ({processed[0]}/{total_blocks})", end="\r", flush=True)

    if total_blocks > 0:
        update_progress(processed, total_blocks)

    current_chunk = []
    current_indices = []
    current_length = 0

    def process_chunk(processed):

        if not current_chunk:
            return

        combined = "\n<<<SPLIT>>>\n".join(current_chunk)
        translated_combined = safe_translate(combined)
        split_texts = translated_combined.split("\n<<<SPLIT>>>\n")

        if len(split_texts) != len(current_chunk):
            split_texts = [safe_translate(t) for t in current_chunk]

        for idx, txt in zip(current_indices, split_texts):
            translated_entries[idx] = txt
            processed[0] += 1
            update_progress(processed, total_blocks)

    # Формування chunk
    for i, (timing, text) in enumerate(subtitle_entries):

        if timing is None:
            translated_entries[i] = text
            continue

        text_length = len(text)

        if current_length + text_length > MAX_CHARS and current_chunk:
            process_chunk(processed)
            current_chunk = []
            current_indices = []
            current_length = 0

        current_chunk.append(text)
        current_indices.append(i)
        current_length += text_length

    # переклад останнього chunk
    process_chunk(processed)
    print()

    # =============================
    # ВІДНОВЛЕННЯ SRT + АВТОНУМЕРАЦІЯ
    # =============================

    translated_blocks = []
    original_full = []
    translated_full = []
    translated_template = []

    counter = 1

    for i, (timing, text) in enumerate(subtitle_entries):

        if timing is None:
            translated_blocks.append(text)
            continue

        translated_text = translated_entries[i]

        new_block = "\n".join([str(counter), timing, translated_text])
        translated_blocks.append(new_block)

        original_full.extend([str(counter), timing, text, ""])
        translated_full.extend([str(counter), timing, translated_text, ""])
        translated_template.extend([str(counter), timing, "", ""])

        counter += 1

    # ==================== ЗАПИС ФАЙЛІВ ========================

    # ==========================================================
    # Підготовка основної теки
    # ==========================================================
    base_folder = output_folder or os.path.dirname(input_path)
    base_name = os.path.splitext(os.path.basename(input_path))[0]
    if output_name:
        filename = f"{output_name}_{base_name}"
    else:
        filename = base_name
    new_folder = os.path.join(base_folder, filename)
    os.makedirs(new_folder, exist_ok=True)

    # ==========================================================
    # SRT
    # ==========================================================
    srt_folder = os.path.join(new_folder, "SRT")
    os.makedirs(srt_folder, exist_ok=True)

    output_srt_path = os.path.join(srt_folder, f"UA_{filename}.srt")

    with open(output_srt_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(translated_blocks))

    # ==========================================================
    # BILINGUAL
    # ==========================================================
    bilingual_doc_path = os.path.join(new_folder, f"BL_{filename}.docx")
    bilingual_doc = Document()

    bilingual_doc.add_heading("Original:", level=1)
    for line in original_full:
        if line.strip().isdigit():
            bilingual_doc.add_heading(line, level=2)
        else:
            bilingual_doc.add_paragraph(line)
    bilingual_doc.add_page_break()

    bilingual_doc.add_heading("Ukraine AutoTranslate:", level=1)
    for line in translated_full:
        if line.strip().isdigit():
            bilingual_doc.add_heading(line, level=2)
        else:
            bilingual_doc.add_paragraph(line)
    bilingual_doc.add_page_break()

    # bilingual_doc.add_heading("MyTranslate:", level=1)
    # for line in translated_template:
    #     if line.strip().isdigit():
    #         bilingual_doc.add_heading(line, level=2)
    #     else:
    #         bilingual_doc.add_paragraph(line)
    # bilingual_doc.save(bilingual_doc_path)

    # ==========================================================
    # FORMATTED
    # ==========================================================
    formatted_folder = os.path.join(new_folder, "FORMATTED")
    os.makedirs(formatted_folder, exist_ok=True)

    # PLAIN TEXT
    plain_doc_path = os.path.join(formatted_folder, f"ORIG_FORM_{filename}.docx")
    plain_doc = Document()
    plain_doc.add_heading("Original:", level=1)

    original_blocks_plain = []
    for timing, text in subtitle_entries:
        if timing:
            cleaned = text.replace("\n", "|")
            original_blocks_plain.append(f"||{cleaned}||")
    plain_doc.add_paragraph("\n".join(original_blocks_plain))
    plain_doc.add_page_break()

    plain_doc.add_heading("Ukraine AutoTranslate:", level=1)
    translated_blocks_plain = []
    for i, (timing, text) in enumerate(subtitle_entries):
        if timing:
            cleaned = translated_entries[i].replace("\n", "|")
            translated_blocks_plain.append(f"||{cleaned}||")
    plain_doc.add_paragraph("\n".join(translated_blocks_plain))
    plain_doc.save(plain_doc_path)

    # AUTHOR TEMPLATE
    author_template_path = os.path.join(formatted_folder, f"MY_FORM_{filename}.docx")
    author_doc = Document()
    author_doc.add_heading("Author Version:", level=1)
    for timing, text in subtitle_entries:
        if timing:
            author_doc.add_paragraph("")
    author_doc.save(author_template_path)

    # ==========================================================
    # SCRIPT
    # ==========================================================
    script_folder = os.path.join(new_folder, "SCRIPT")
    os.makedirs(script_folder, exist_ok=True)

    original_script_path = os.path.join(script_folder, f"ORIG_SUB_{filename}.docx")
    translated_script_path = os.path.join(script_folder, f"UA_SUB_{filename}.docx")

    original_script_doc = Document()
    translated_script_doc = Document()

    # original_script_doc.add_heading("Original Script (Clean Dialogue)", level=1)
    # translated_script_doc.add_heading("Ukraine AutoTranslate Script (Clean Dialogue)", level=1)

    for i, (timing, text) in enumerate(subtitle_entries):
        if timing is None:
            continue
        clean_original = text.replace("\n", " ").strip()
        if clean_original:
            original_script_doc.add_paragraph(clean_original)
        clean_translated = translated_entries[i].replace("\n", " ").strip()
        if clean_translated:
            translated_script_doc.add_paragraph(clean_translated)

    original_script_doc.save(original_script_path)
    translated_script_doc.save(translated_script_path)

    # ==========================================================
    # RETURN RESULT (ВАЖЛИВО ДЛЯ GUI)
    # ==========================================================

    return {
        "folder": new_folder,
        "srt": output_srt_path,
        "bilingual": bilingual_doc_path,
        "plain": plain_doc_path,
        "author_template": author_template_path,
        "orig_script": original_script_path,
        "translated_script": translated_script_path,
        "cache": CACHE_FILE
    }

    # ==========================================================
    # ФІНАЛ
    # ==========================================================
    # print("\n === Готово! ===\n")
    # print(" Папка створена:", new_folder)
    # print(" Створено файли:")
    # print(" -", output_srt_path)
    # print(" -", bilingual_doc_path)
    # print(" -", plain_doc_path)
    # print(" -", author_template_path)
    # print(" -", original_script_path)
    # print(" -", translated_script_path)
    # print(" -", CACHE_FILE, "(глобальний кеш)\n")

# ==========================================================
# =====================  MAIN  ============================
# ==========================================================

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Використання: python translate.py <шлях_до_srt/ass>")
        sys.exit(1)
    
    # Шлях до файлу береться з аргументів командного рядка
    input_file = sys.argv[1]

    # Викликаємо функцію перекладу
    run_translate(input_file)